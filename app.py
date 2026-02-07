import os
import json
import sys
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from flask_socketio import SocketIO, emit, join_room

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

# Determine the folder where the executable (or script) is located
# This is where we will store the 'data' folder
if getattr(sys, 'frozen', False):
    # If the application is run as a bundle (PyInstaller)
    BASE_DIR = os.path.dirname(sys.executable)
    # Static assets are inside the MEIPASS folder
    STATIC_FOLDER = resource_path('.')
else:
    # If run from Python script
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    STATIC_FOLDER = '.'

app = Flask(__name__, static_url_path='', static_folder=STATIC_FOLDER)
CORS(app)  # Enable Cross-Origin Resource Sharing
socketio = SocketIO(app, cors_allowed_origins="*")

# Configuration
PORT = 3000
DATA_DIR = os.path.join(BASE_DIR, 'data')

# Ensure data directory exists
if not os.path.exists(DATA_DIR):
    os.makedirs(DATA_DIR)

# Presence Tracking
socket_to_user = {} # sid -> user_id
user_sockets = {} # user_id -> set(sids)

@app.route('/')
def serve_index():
    return send_from_directory(app.static_folder, 'index.html')

@app.route('/<path:path>')
def serve_static(path):
    return send_from_directory(app.static_folder, path)

# WebSocket Events
@socketio.on('join')
def on_join(data):
    user_id = data.get('userId')
    sid = request.sid
    
    if user_id:
        join_room(user_id)
        
        # Update mappings
        socket_to_user[sid] = user_id
        if user_id not in user_sockets:
            user_sockets[user_id] = set()
        user_sockets[user_id].add(sid)
        
        count = len(user_sockets[user_id])
        print(f"Client {sid} joined room {user_id}. Total count: {count}")
        
        emit('status', {'msg': f'Joined room {user_id}'}, room=user_id)
        
        # If multiple users are present, warn everyone
        if count > 1:
            emit('presence_warning', {
                'message': '目前已有其他用戶修改內容，避免資料無法同步，請稍等',
                'count': count
            }, room=user_id)

@socketio.on('disconnect')
def on_disconnect():
    sid = request.sid
    user_id = socket_to_user.get(sid)
    
    if user_id:
        if user_id in user_sockets:
            user_sockets[user_id].discard(sid)
            if len(user_sockets[user_id]) == 0:
                del user_sockets[user_id]
        
        if sid in socket_to_user:
            del socket_to_user[sid]
            
        print(f"Client {sid} disconnected from {user_id}. Remaining: {len(user_sockets.get(user_id, []))}")

# API: Login
@app.route('/api/login', methods=['POST'])
def login():
    data = request.json
    user_id = data.get('userId')

    if not user_id:
        return jsonify({'success': False, 'message': 'User ID is required'}), 400

    # In a real app, you'd check passwords here.
    # We allow "Spe for u" or any ID for this local version.
    return jsonify({'success': True, 'message': 'Login successful'})

# API: Get Data
@app.route('/api/data/<user_id>', methods=['GET'])
def get_data(user_id):
    file_path = os.path.join(DATA_DIR, f"{user_id}.json")
    
    try:
        if os.path.exists(file_path):
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return jsonify({'success': True, 'data': data})
        else:
            return jsonify({'success': True, 'data': None})
    except Exception as e:
        print(f"Error reading data: {e}")
        return jsonify({'success': False, 'message': 'Internal Server Error'}), 500

# API: Save Data
@app.route('/api/data/<user_id>', methods=['POST'])
def save_data(user_id):
    req_data = request.json
    
    # Check if this is a legacy request (direct data) or enveloped request ({data: ..., lastSyncedTimestamp: ...})
    if 'data' in req_data and 'lastSyncedTimestamp' in req_data:
        new_data = req_data['data']
        client_timestamp = req_data['lastSyncedTimestamp']
        force_save = req_data.get('force', False)
    else:
        # Backward compatibility or simple save
        new_data = req_data
        client_timestamp = None
        force_save = True # Assume force for legacy calls to avoid breaking changes immediately

    file_path = os.path.join(DATA_DIR, f"{user_id}.json")
    
    try:
        # Check for conflicts if not forcing
        if not force_save and os.path.exists(file_path):
            with open(file_path, 'r', encoding='utf-8') as f:
                existing_file_content = json.load(f)
            
            # Extract timestamp from existing data
            existing_timestamp = existing_file_content.get('timestamp')
            
            if str(existing_timestamp) != str(client_timestamp):
                print(f"Conflict detected for {user_id}. Server: {existing_timestamp}, Client saw: {client_timestamp}")
                return jsonify({
                    'success': False, 
                    'message': 'Data conflict detected. Please reload.',
                    'serverData': existing_file_content
                }), 409

        # Write new data
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(new_data, f, ensure_ascii=False, indent=2)
        
            # Broadcast update to room
        try:
            # We broadcast the new timestamp so clients know there is a new version
            # We don't broadcast full data (efficiency), just notification
            socket_id = req_data.get('socketId') # Get sender's socket ID
            socketio.emit('data_updated', {
                'timestamp': new_data.get('timestamp'),
                'sourceSocketId': socket_id, # Echo back the sender's socket ID
                'updater': request.remote_addr  # Optional: who updated?
            }, room=user_id)
            print(f"Broadcasted update for room {user_id}")
        except Exception as e:
            print(f"Socket emit failed: {e}")

        return jsonify({'success': True, 'message': 'Data saved successfully'})
    except Exception as e:
        print(f"Error writing data: {e}")
        return jsonify({'success': False, 'message': 'Internal Server Error'}), 500

# --- WORD EXPORT FUNCTIONALITY ---
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from flask import send_file

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    # Ensure explicitly defined
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes matters
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_cell_background(cell, color_hex):
    """
    Set cell background color
    color_hex: e.g., "F2F9FF" (without #)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def generate_word_schedule(data):
    document = Document()
    
    # Page Setup (A4 Portrait)
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)

    # Title
    title_text = data.get('title', '特教班課表')
    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title_text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = '標楷體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

    # Date (Right aligned)
    # Using specific font for print match if possible, but standard is fine


    # Table Setup
    # Columns: Fri, Thu, Wed, Tue, Mon, Time, Section (7 columns)
    table = document.add_table(rows=0, cols=7)
    table.style = 'Table Grid'
    table.autofit = False
    
    # Column Widths (Approximation based on visual: Period/Time narrow, Days equal)
    # Total width approx 18.46cm (21 - 1.27*2)
    # Ratios: Days ~16%, Time ~12%, Section ~8%
    # 5 days * 16 = 80, + 12 + 8 = 100
    
    # Header Row
    hdr_cells = table.add_row().cells
    headers = ["星期五", "星期四", "星期三", "星期二", "星期一", "時間", "節次"]
    widths = [Cm(2.9), Cm(2.9), Cm(2.9), Cm(2.9), Cm(2.9), Cm(2.5), Cm(1.5)] # Total ~18.5
    
    for i, text in enumerate(headers):
        cell = hdr_cells[i]
        cell.width = widths[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = '標楷體'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        # Center vertical
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)

    # Data Rows
    # Expected schedule_rows format: list of dicts { 'period': '1', 'name': '1', 'time': '...', 'days': { 'monday': '...', ... }, 'isLunch': bool }
    schedule_rows = data.get('schedule', [])

    for row_data in schedule_rows:
        row = table.add_row()
        cells = row.cells
        
        # Row Height
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '800') # ~1.4cm min height
        if row_data.get('isLunch'):
             trHeight.set(qn('w:val'), '700') 
        trPr.append(trHeight)

        if row_data.get('isLunch'):
            # Merge first 5 cells for "Lunch"
            # Logic: cells[0] is Fri, cells[4] is Mon. Merge cells[0] to cells[4].
            # docx table merge: cell_1.merge(cell_2)
            
            merged_cell = cells[0].merge(cells[4])
            merged_cell.text = "午休"
            p = merged_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(14)
            run.font.name = '標楷體'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
            
            # Time Cell
            cells[5].text = "12:30\n|\n13:10"
            p = cells[5].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Section Cell (Empty)
            cells[6].text = ""
            
        else:
            # Standard Row
            # Map frontend days to columns: Fri(0), Thu(1), Wed(2), Tue(3), Mon(4)
            day_map = {0: 'friday', 1: 'thursday', 2: 'wednesday', 3: 'tuesday', 4: 'monday'}
            
            for i in range(5):
                day_key = day_map[i]
                cell_content = row_data['days'].get(day_key, '')
                # Handle newlines for multi-line content (courses)
                # Frontend sends <br> or \n? Let's assume \n or clean up.
                cell = cells[i]
                # Replace <br> with \n just in case
                clean_content = cell_content.replace('<br>', '\n').replace('&nbsp;', ' ')
                cell.text = clean_content
                
                # Formatting
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if p.runs:
                    run = p.runs[0]
                    run.font.size = Pt(12) # Content size
                    run.font.name = '標楷體'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

            # Time
            time_str = row_data.get('time', '').replace('~', '\n|\n')
            cells[5].text = time_str
            p = cells[5].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Section Name
            cells[6].text = row_data.get('name', '')
            p = cells[6].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Vertical Align Center for all cells in row
        for cell in cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)

    # Date (Right aligned) - Moved to bottom
    date_text = data.get('date_range', '')
    if date_text:
        date_p = document.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Add some space before date?
        date_p.paragraph_format.space_before = Pt(12) 
        date_run = date_p.add_run(f"實施日期 {date_text}")
        date_run.font.size = Pt(12)
        date_run.font.name = 'Times New Roman'
        date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

    # Save to buffer
    f = BytesIO()
    document.save(f)
    f.seek(0)
    return f

def generate_word_teacher_schedule(data):
    document = Document()
    
    # Page Setup (A4 Portrait)
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)

    teachers_data = data.get('teachers', [])
    
    for idx, teacher in enumerate(teachers_data):
        if idx > 0:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # Title
        title_text = data.get('title', '特教班教師課表')
        title_p = document.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(20)
        run.font.name = '標楷體'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

        # Subheader: Teacher Name
        name_p = document.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_run = name_p.add_run(f"任課教師：{teacher['name']} 老師")
        name_run.font.size = Pt(16)
        name_run.font.name = '標楷體'
        name_run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        
        # Table
        table = document.add_table(rows=0, cols=7)
        table.style = 'Table Grid'
        table.autofit = False
        
        # New Requirement: Time=1.5, Per=1.5
        # Remaining width approx 18.5 - 3.0 = 15.5cm. 15.5/5 = 3.1cm per day
        widths = [Cm(3.1), Cm(3.1), Cm(3.1), Cm(3.1), Cm(3.1), Cm(1.5), Cm(1.5)]
        
        hdr_cells = table.add_row().cells
        headers = ["星期五", "星期四", "星期三", "星期二", "星期一", "時間", "節次"]
        for i, text in enumerate(headers):
            cell = hdr_cells[i]
            cell.width = widths[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = '標楷體'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)

        rows = teacher.get('schedule_rows', [])
        for row_data in rows:
            row = table.add_row()
            cells = row.cells
            
            # Enforce column widths for every row
            for ck, width in enumerate(widths):
                cells[ck].width = width
            
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), '850')
            if row_data.get('isLunch'): trHeight.set(qn('w:val'), '700')
            trPr.append(trHeight)

            if row_data.get('isLunch'):
                merged_cell = cells[0].merge(cells[4])
                merged_cell.text = "午休"
                p = merged_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.runs[0]
                run.font.size = Pt(14)
                run.font.name = '標楷體'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                
                cells[5].text = "12:30\n|\n13:10"
                p = cells[5].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[6].text = ""
            else:
                day_map = {0: 'friday', 1: 'thursday', 2: 'wednesday', 3: 'tuesday', 4: 'monday'}
                for i in range(5):
                    day_key = day_map[i]
                    content = row_data['days'].get(day_key, '')
                    content = content.replace('<br>', '\n').replace('&nbsp;', ' ')
                    # Remove HTML tags if any (basic removal)
                    if '<' in content:
                        import re
                        content = re.sub('<[^<]+?>', '', content)
                        
                    cell = cells[i]
                    cell.text = content
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if p.runs:
                        run = p.runs[0]
                        run.font.size = Pt(12) # Adjusted size
                        run.font.name = '標楷體'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                
                time_str = row_data.get('time', '').replace('~', '\n|\n')
                cells[5].text = time_str
                p = cells[5].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                cells[6].text = row_data.get('name', '')
                p = cells[6].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for cell in cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)

        # Footer Date
        date_text = data.get('date_range', '')
        if date_text:
            date_p = document.add_paragraph()
            date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_p.paragraph_format.space_before = Pt(6)
            run = date_p.add_run(f"實施日期 {date_text}")
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

        # Footer Stats Text
        stats_text = teacher.get('stats_text', '')
        if stats_text:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(stats_text)
            run.font.size = Pt(12)
            run.font.name = '標楷體'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

        # Footer Stats Table (Total, Base, PartTime, Overtime)
        stats = teacher.get('stats_table', {})
        st_table = document.add_table(rows=1, cols=4)
        st_table.style = 'Table Grid'
        st_table.allow_autofit = True
        
        row = st_table.rows[0]
        items = [
            f"總時數：{stats.get('total', 0)} 節",
            f"基本鐘點：{stats.get('base', 0)} 節 {stats.get('note', '')}",
            f"兼課：{stats.get('part_time', 0)} 節",
            f"超鐘點：{stats.get('overtime', '')}"
        ]
        
        for i, text in enumerate(items):
            cell = row.cells[i]
            
            # 特別處理"基本鐘點"欄位，讓備註換行
            if i == 1 and stats.get('note', ''):  # i==1 是基本鐘點欄位
                cell.text = ""  # 先清空
                # 第一行：基本鐘點數
                p = cell.paragraphs[0]
                run = p.add_run(f"基本鐘點：{stats.get('base', 0)} 節")
                run.font.size = Pt(12)
                run.font.name = '標楷體'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                
                # 第二行：備註
                p2 = cell.add_paragraph()
                run2 = p2.add_run(stats.get('note', ''))
                run2.font.size = Pt(12)
                run2.font.name = '標楷體'
                run2._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
            else:
                cell.text = text
                p = cell.paragraphs[0]
                run = p.runs[0]
                run.font.size = Pt(12)
                run.font.name = '標楷體'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')


    f = BytesIO()
    document.save(f)
    f.seek(0)
    return f



@app.route('/api/export/word/teacher', methods=['POST'])
def export_word_teacher():
    try:
        data = request.json
        print("Generating Teacher Word...")
        file_stream = generate_word_teacher_schedule(data)
        return send_file(file_stream, as_attachment=True, download_name='teachers_schedule.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500







@app.route('/api/export/word', methods=['POST'])
def export_word():
    try:
        data = request.json
        print("Generating Word document...")
        
        file_stream = generate_word_schedule(data)
        
        return send_file(
            file_stream,
            as_attachment=True,
            download_name='schedule.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        print(f"Error generating Word: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/export/word/student', methods=['POST'])
def export_word_student():
    try:
        data = request.json
        print("Generating Student Word...")
        file_stream = generate_word_student_schedule(data)
        return send_file(
            file_stream, 
            as_attachment=True, 
            download_name='student_schedules.docx', 
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500

def generate_word_student_schedule(data):
    document = Document()
    
    # Page Setup (A4 Portrait)
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)

    students_data = data.get('students', [])
    title_text = data.get('title', '學生課表')
    
    for idx, student in enumerate(students_data):
        if idx > 0:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # Header Title
        title_p = document.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(20)
        run.font.name = '標楷體'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

        # Student Name
        name_p = document.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_run = name_p.add_run(f"{student['name']}")
        name_run.font.size = Pt(14)
        name_run.bold = True
        name_run.font.name = '標楷體'
        name_run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        
        # Table
        # Cols: Fri, Thu, Wed, Tue, Mon, Time, Per
        table = document.add_table(rows=0, cols=7)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Original Widths: [Cm(2.95), Cm(2.95), Cm(2.95), Cm(2.95), Cm(2.95), Cm(2.2), Cm(1.5)]
        # New Requirement: Time=1.5, Per=1.5
        # Remaining width approx 18.5 - 3.0 = 15.5cm. 15.5/5 = 3.1cm per day
        widths = [Cm(3.1), Cm(3.1), Cm(3.1), Cm(3.1), Cm(3.1), Cm(1.5), Cm(1.5)]
        
        hdr_cells = table.add_row().cells
        headers = ["星期五", "星期四", "星期三", "星期二", "星期一", "時間", "節次"]
        for i, text in enumerate(headers):
            cell = hdr_cells[i]
            cell.width = widths[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = '標楷體'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)

        rows = student.get('schedule_rows', [])
        for row_data in rows:
            row = table.add_row()
            cells = row.cells
            
            # Enforce column widths for every row
            for ck, width in enumerate(widths):
                cells[ck].width = width

            # Row Height
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), '1300') # Matches approx 50-60px height
            trPr.append(trHeight)

            day_map = {0: 'friday', 1: 'thursday', 2: 'wednesday', 3: 'tuesday', 4: 'monday'}
            for i in range(5):
                day_key = day_map[i]
                content = row_data['days'].get(day_key, '')
                content = content.replace('<br>', '\n').replace('&nbsp;', ' ')
                
                cell = cells[i]
                # Clear existing paragraphs
                cell._element.clear_content()
                
                # Check previous logic for split usage
                if content.strip():
                    parts = content.split('\n')
                    
                    # 1. Subject (Large, Bold)
                    if len(parts) > 0 and parts[0].strip():
                        p = cell.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(parts[0].strip())
                        run.bold = True
                        run.font.size = Pt(16)
                        run.font.name = '標楷體'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                    
                    # 2. Teacher (Small)
                    if len(parts) > 1 and parts[1].strip():
                        p = cell.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(parts[1].strip())
                        run.font.size = Pt(10)
                        run.font.name = '標楷體'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

                    # 3. Room (Small)
                    if len(parts) > 2 and parts[2].strip():
                        p = cell.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(parts[2].strip())
                        run.font.size = Pt(10)
                        run.font.name = '標楷體'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                else:
                    # Empty cell
                    cell.add_paragraph('')

            # Time
            time_raw = row_data.get('time', '')
            if '~' in time_raw:
                t_parts = time_raw.split('~')
                cells[5].text = "" # Clear
                p = cells[5].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Start Time
                run = p.add_run(t_parts[0])
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
                
                # Pipe
                p.add_run('\n|\n')
                
                # End Time
                run = p.add_run(t_parts[1])
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
            else:
                cells[5].text = time_raw
                p = cells[5].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Section
            cells[6].text = row_data.get('name', '')
            p = cells[6].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].font.size = Pt(12)
            
            for cell in cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)

    f = BytesIO()
    document.save(f)
    f.seek(0)
    return f

if __name__ == '__main__':
    print(f"Server is running at http://localhost:{PORT}")
    print(f"To share with other computers, use your IP address, e.g., http://192.168.x.x:{PORT}")
    # app.run(host='0.0.0.0', port=PORT, debug=True)
    socketio.run(app, host='0.0.0.0', port=PORT, debug=True, allow_unsafe_werkzeug=True)
