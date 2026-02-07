import unittest
import os
import sys
from docx import Document
from docx.shared import Cm
from io import BytesIO

# Import app.py (assuming it's in the same directory)
try:
    from app import generate_word_schedule, generate_word_teacher_schedule, generate_word_student_schedule
except ImportError:
    sys.path.append(os.path.dirname(os.path.abspath(__file__)))
    from app import generate_word_schedule, generate_word_teacher_schedule, generate_word_student_schedule

class TestWordExport(unittest.TestCase):

    def setUp(self):
        # Mock Data for General Schedule
        self.mock_schedule_data = {
            'title': '測試課表',
            'date_range': '113/02/01 ~ 113/06/30',
            'schedule': [
                {
                    'period': '1', 'name': '第一節', 'time': '08:00~08:50',
                    'days': {'monday': '國文\n王老師', 'tuesday': '數學\n李老師', 'wednesday': '', 'thursday': '', 'friday': ''},
                    'isLunch': False
                },
                {
                    'period': 'Lunch', 'name': '午休', 'time': '12:00~13:00',
                    'days': {},
                    'isLunch': True
                }
            ]
        }

        # Mock Data for Teacher Schedule
        self.mock_teacher_data = {
            'title': '測試教師課表',
            'date_range': '113/02/01 ~ 113/06/30',
            'teachers': [
                {
                    'name': '王老師',
                    'stats_text': '總時數：10',
                    'stats_table': {'total': 10, 'base': 10, 'part_time': 0, 'overtime': 0, 'note': '(每週需搭車往返中心上課)'},
                    'schedule_rows': [
                        {
                            'period': '1', 'time': '08:00~08:50',
                            'days': {'monday': '國文', 'tuesday': '', 'wednesday': '', 'thursday': '', 'friday': ''},
                            'isLunch': False
                        }
                    ]
                }
            ]
        }

        # Mock Data for Student Schedule
        self.mock_student_data = {
            'title': '測試學生課表',
            'students': [
                {
                    'name': '張三',
                    'schedule_rows': [
                        {
                            'period': '1', 'time': '08:00~08:50',
                            'days': {'monday': '國文\n王老師\n101', 'tuesday': '', 'wednesday': '', 'thursday': '', 'friday': ''},
                            'isLunch': False
                        }
                    ]
                }
            ]
        }

    def verify_column_widths(self, table, expected_widths, tolerance=0.1):
        """Helper to verify column widths of the first row"""
        rows = table.rows
        if not rows:
            return False
            
        cells = rows[0].cells
        if len(cells) != len(expected_widths):
            print(f"Column count mismatch: got {len(cells)}, expected {len(expected_widths)}")
            return False

        for i, cell in enumerate(cells):
            # cell.width is in EMUs or None if not set? 
            # python-docx cell.width returns Emu. 1 Cm = 360000 Emu
            # But sometimes it's None if autofit. The code explicitly sets it.
            
            # Note: Word internal units are Twips? python-docx uses Emu.
            # 1 Cm = 360000 Emu
            if cell.width is None:
                # If width isn't set on the cell directly (sometimes it's on the column), this might be None
                # But generated code sets cell.width
                print(f"Cell {i} width is None")
                continue

            width_cm = cell.width.cm
            expected = expected_widths[i]
            
            if abs(width_cm - expected) > tolerance:
                print(f"Column {i} width mismatch: got {width_cm}, expected {expected}")
                return False
        return True

    def test_schedule_export_format(self):
        """簡易課表匯出測試"""
        # Generate
        stream = generate_word_schedule(self.mock_schedule_data)
        doc = Document(stream)
        
        # 1. Verify Page Size (A4)
        section = doc.sections[0]
        self.assertAlmostEqual(section.page_width.cm, 21.0, delta=0.1)
        self.assertAlmostEqual(section.page_height.cm, 29.7, delta=0.1)

        # 2. Verify Table Exists
        self.assertTrue(len(doc.tables) > 0)
        table = doc.tables[0]
        
        # 3. Verify Columns (7 cols)
        self.assertEqual(len(table.columns), 7)

        # 4. Verify Widths (Fri~Mon: 2.9, Time: 2.5, Per: 1.5)
        # Note: The Loop in app.py sets headers: Fri, Thu, Wed, Tue, Mon, Time, Per
        # Widths: [2.9, 2.9, 2.9, 2.9, 2.9, 2.5, 1.5]
        expected_widths = [2.9, 2.9, 2.9, 2.9, 2.9, 2.5, 1.5]
        self.assertTrue(self.verify_column_widths(table, expected_widths))

        # 5. Verify Lunch Merge (Row 2 in mock data -> Row 2 in table (Header is Row 0))
        # Wait, data loop starts adding rows. Header is row 0.
        # Data index 0 -> Row 1. Data index 1 (Lunch) -> Row 2.
        lunch_row = table.rows[2]
        cells = lunch_row.cells
        # Merged cells share the same object in python-docx? Or text is same?
        # Typically mechanism: cell[0] is the merged cell. cell[1]...cell[4] should be "part of" it?
        # In python-docx, if merged, accessing cells[1] might return the same cell object as cells[0] 
        # OR different objects but checking .text confirms content availability.
        
        self.assertEqual(cells[0].text, "午休")
        # Check if cells[4] (Monday) is merged with cells[0] (Friday)
        # One way is to check grid_span if exposed, but python-docx read support is limited.
        # User requirement: "isLunch: true implies merge".
        # We can check the text of the first cell is "午休"

    def test_teacher_export_format(self):
        """教師課表匯出測試"""
        stream = generate_word_teacher_schedule(self.mock_teacher_data)
        doc = Document(stream)
        
        table = doc.tables[0]
        
        # Widths: [3.1, 3.1, 3.1, 3.1, 3.1, 1.5, 1.5]
        expected_widths = [3.1, 3.1, 3.1, 3.1, 3.1, 1.5, 1.5]
        self.assertTrue(self.verify_column_widths(table, expected_widths))

        # Check Teacher Name in doc
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        self.assertIn("王老師", full_text)

    def test_student_export_format(self):
        """學生課表匯出測試"""
        stream = generate_word_student_schedule(self.mock_student_data)
        doc = Document(stream)
        
        table = doc.tables[0]
        
        # Widths: [3.1, 3.1, 3.1, 3.1, 3.1, 1.5, 1.5]
        expected_widths = [3.1, 3.1, 3.1, 3.1, 3.1, 1.5, 1.5]
        self.assertTrue(self.verify_column_widths(table, expected_widths))
        
        # Check Content Format
        # Row 1 (Index 1) -> Monday (Index 4)
        row1 = table.rows[1]
        cell_mon = row1.cells[4]
        # Text should be concatenated
        self.assertIn("國文", cell_mon.text)
        self.assertIn("王老師", cell_mon.text)
        
        # Check Row Height
        # python-docx doesn't easily expose trHeight reading without Oxml inspection.
        # We can skip exact height validation or use Oxml if critical.
        # For now, let's assume if function runs without error using the code we saw, it works.

    def test_teacher_note_newline(self):
        """測試教師課表基本鐘點備註換行功能"""
        stream = generate_word_teacher_schedule(self.mock_teacher_data)
        doc = Document(stream)
        
        # 找到統計資料表格（應該是第二個表格，第一個是課表）
        self.assertTrue(len(doc.tables) >= 2, "應該至少有兩個表格（課表和統計）")
        stats_table = doc.tables[1]
        
        # 統計表格應該有4列（總時數、基本鐘點、兼課、超鐘點）
        row = stats_table.rows[0]
        self.assertEqual(len(row.cells), 4, "統計表格應該有4欄")
        
        # 基本鐘點在第二欄（index 1）
        base_cell = row.cells[1]
        
        # 檢查這個儲存格有多個段落（表示有換行）
        paragraphs = base_cell.paragraphs
        self.assertGreaterEqual(len(paragraphs), 2, "基本鐘點儲存格應該至少有兩個段落（鐘點數和備註）")
        
        # 第一段應該包含基本鐘點數
        self.assertIn("基本鐘點", paragraphs[0].text)
        self.assertIn("10", paragraphs[0].text)
        self.assertIn("節", paragraphs[0].text)
        
        # 第二段應該只包含備註
        self.assertEqual(paragraphs[1].text.strip(), "(每週需搭車往返中心上課)")


if __name__ == '__main__':
    unittest.main()
